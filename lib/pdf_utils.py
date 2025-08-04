from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt

def update_resume_objective(
    source_path: str,
    output_path: str,
    new_objective: str,
    placeholder: str = '<objective_here>',
    font_name: str = 'Spectral',
    font_size: int = 10
) -> str:
    """
    Update the objective section in a DOCX file.
    
    Args:
        source_path: Path to the source DOCX file
        output_path: Path where to save the updated file
        new_objective: The new objective text to insert
        placeholder: The placeholder text to replace
        font_name: Font name for the objective text
        font_size: Font size for the objective text
        
    Returns:
        Path to the saved file
        
    Raises:
        ValueError: If the placeholder is not found in the document
    """
    doc = Document(source_path)
    found_placeholder = False
        
    # Search through all paragraphs in the document
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Clear the paragraph and add a new run with the new objective
            paragraph.clear()
            run = paragraph.add_run(new_objective)
            font = run.font
            font.name = font_name
            font.size = Pt(font_size)
            found_placeholder = True
            break
    
    if not found_placeholder:
        raise ValueError(f"Placeholder text '{placeholder}' not found in the document")
    
    # Save the updated document
    doc.save(output_path)
    return output_path
    
    # Save the updated document
    # output_path = str(Path(output_path).with_suffix('.docx'))
    # doc.save(output_path)
    # return output_path

def create_pdf_from_docx(docx_path: str, output_dir: Optional[str] = None) -> str:
    """
    Convert a DOCX file to PDF using LibreOffice.
    
    Args:
        docx_path: Path to the DOCX file
        output_dir: Directory to save the PDF (defaults to same as input)
        
    Returns:
        Path to the generated PDF file
    """
    from utils.pdf_utils import convert_to_pdf_libreoffice
    
    if output_dir is None:
        output_dir = str(Path(docx_path).parent)

    if not Path(output_dir).exists():
        Path(output_dir).mkdir(parents=True, exist_ok=True)

    print('output_dir', output_dir)
    
    pdf_path = Path(output_dir)
    pdf_name = f"{Path(docx_path).stem}.pdf"
    convert_to_pdf_libreoffice(docx_path, str(pdf_path), pdf_name)
    return str(pdf_path / pdf_name)
