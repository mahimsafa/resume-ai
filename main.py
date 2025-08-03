#!/usr/bin/env python3
"""
Resume AI - AI-powered resume objective generator

This script generates a tailored career objective based on your resume and a job description,
then updates a DOCX resume with the new objective.

Usage:
    python main.py [job_description_file]

Example:
    python main.py input/jobdescription.txt
"""
import os
import sys
import re
import argparse
from pathlib import Path
from typing import Optional, Tuple, Dict, Any

# Third-party imports
from langchain_google_vertexai import ChatVertexAI
from langchain.prompts import PromptTemplate
from docx import Document
from docx.shared import Pt
from docx.document import Document as DocumentType

# Local imports
from utils import (
    ensure_extension,
    get_unique_filename,
    clean_text_for_filename,
    ensure_directory,
    read_file as read_file_util
)

class ResumeAI:
    """Main application class for Resume AI."""
    
    def __init__(self, base_dir: Optional[str] = None):
        """Initialize the Resume AI application.
        
        Args:
            base_dir: Base directory for the application (defaults to current directory)
        """
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()
        self.input_dir = self.base_dir / 'input'
        self.output_dir = self.base_dir / 'generated'
        self.llm = ChatVertexAI(model="gemini-2.5-flash", temperature=0.7)
        
        # Ensure required directories exist
        self._setup_directories()
    
    def _setup_directories(self) -> None:
        """Ensure required directories exist."""
        ensure_directory(self.input_dir)
        ensure_directory(self.output_dir)
    
    def _check_required_files(self) -> Tuple[bool, str]:
        """Check if required input files exist.
        
        Returns:
            Tuple of (success, message)
        """
        required_files = {
            'resume': self.input_dir / 'resume.md',
            'job_description': self.input_dir / 'jobdescription.txt',
            'resume_template': self.input_dir / 'resume.docx'
        }
        
        missing = [name for name, path in required_files.items() if not path.exists()]
        if missing:
            return False, f"Missing required files: {', '.join(missing)}. Please add them to the 'input' directory."
        return True, ""
    
    def _generate_objective_and_filename(
        self,
        resume_content: str,
        job_description: str,
        base_name: str = "resume"
    ) -> Tuple[str, str]:
        """Generate a tailored objective and filename based on resume and job description.
        
        Args:
            resume_content: The content of the resume
            job_description: The job description
            base_name: Base name to use for the filename
            
        Returns:
            Tuple containing (objective, filename)
        """
        prompt_template = """
        Based on the following resume and job description:
        
        RESUME:
        {resume}
        
        JOB DESCRIPTION:
        {job_description}
        
        Please provide:
        1. A tailored career objective (2-3 sentences) that highlights the most relevant skills and experiences.
        2. A suggested filename that includes the job role and company name in the format: [role]-at-[company]
        
        Return the result in this exact format:
        
        OBJECTIVE:
        [Your tailored objective here]
        
        FILENAME:
        [suggested-filename]
        """
        
        prompt = PromptTemplate(
            template=prompt_template,
            input_variables=["resume", "job_description"]
        )
        
        # Format the prompt with the resume and job description
        formatted_prompt = prompt.format(
            resume=resume_content[:4000],  # Limit length to avoid token limits
            job_description=job_description[:2000]
        )
        
        # Generate the response
        response = self.llm.invoke(formatted_prompt)
        
        # Parse the response
        try:
            objective = re.search(r'OBJECTIVE:(.*?)(?=FILENAME:|$)', response.content, re.DOTALL)
            filename = re.search(r'FILENAME:(.*?)$', response.content, re.DOTALL)
            
            if not objective or not filename:
                raise ValueError("Could not parse AI response")
                
            objective = objective.group(1).strip()
            filename = filename.group(1).strip()
            
            # Clean and validate the filename
            filename = clean_text_for_filename(filename)
            if not filename:
                filename = f"{base_name}-generated"
            
            return objective, filename
            
        except Exception as e:
            print(f"Error parsing AI response: {e}")
            # Fallback to a default filename if parsing fails
            return "A highly motivated professional with relevant experience.", f"{base_name}-generated"
    
    def _update_docx_objective(
        self,
        source_path: str,
        output_path: str,
        new_objective: str,
        placeholder: str = '<objective_here>',
        font_name: str = 'Spectral',
        font_size: int = 10
    ) -> str:
        """
        Update the objective section in a DOCX file.
        
        Args:
            source_path: Path to the source DOCX file
            output_path: Path where to save the updated file
            new_objective: The new objective text to insert
            placeholder: The placeholder text to replace
            font_name: Font name for the objective text
            font_size: Font size for the objective text
            
        Returns:
            Path to the saved file
            
        Raises:
            ValueError: If the placeholder is not found in the document
        """
        doc = Document(source_path)
        found_placeholder = False
        
        # Search through all paragraphs in the document
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Clear the paragraph and add a new run with the new objective
                paragraph.clear()
                run = paragraph.add_run(new_objective)
                font = run.font
                font.name = font_name
                font.size = Pt(font_size)
                found_placeholder = True
                break
        
        if not found_placeholder:
            raise ValueError(f"Placeholder text '{placeholder}' not found in the document")
        
        # Save the updated document
        doc.save(output_path)
        return output_path
    
    def run(self, job_description_path: Optional[str] = None):
        """Run the Resume AI application.
        
        Args:
            job_description_path: Optional path to a job description file.
                                 If not provided, uses the default location.
        """
        print("\n=== Resume AI ===\n")
        
        # Set up paths
        if job_description_path:
            job_description_path = Path(job_description_path)
            if not job_description_path.is_absolute():
                job_description_path = self.base_dir / job_description_path
        else:
            job_description_path = self.input_dir / 'jobdescription.txt'
        
        resume_path = self.input_dir / 'resume.md'
        template_path = self.input_dir / 'resume.docx'
        
        # Check if required files exist
        if not job_description_path.exists():
            print(f"Error: Job description file not found at {job_description_path}")
            return
            
        if not resume_path.exists():
            print(f"Error: Resume file not found at {resume_path}")
            return
            
        if not template_path.exists():
            print(f"Error: Resume template (DOCX) not found at {template_path}")
            return
        
        try:
            # Read the resume and job description
            print("Reading resume and job description...")
            resume_content = read_file_util(resume_path)
            job_description = read_file_util(job_description_path)
            
            # Generate the objective and filename
            print("Generating tailored objective...")
            objective, filename = self._generate_objective_and_filename(
                resume_content, 
                job_description
            )
            
            print("\nGenerated Objective:")
            print("-" * 50)
            print(objective)
            print("-" * 50)
            
            # Generate output path
            output_filename = f"{filename}.docx"
            output_path = self.output_dir / output_filename
            
            # Update the DOCX with the new objective
            print(f"\nUpdating resume: {output_path}")
            self._update_docx_objective(
                source_path=template_path,
                output_path=output_path,
                new_objective=objective
            )
            
            print("\n✓ Resume updated successfully!")
            print(f"   Saved to: {output_path}")
            
        except Exception as e:
            print(f"\nError: {str(e)}")
            sys.exit(1)

def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description='Generate a tailored resume objective.')
    parser.add_argument(
        'job_description',
        nargs='?',
        default=None,
        help='Path to the job description file (default: input/jobdescription.txt)'
    )
    return parser.parse_args()

def main():
    """Main entry point for the Resume AI application."""
    args = parse_arguments()
    app = ResumeAI()
    app.run(args.job_description)

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n❌ Operation cancelled by user.")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ An error occurred: {str(e)}", file=sys.stderr)
        print("   Make sure to install the required dependencies:")
        print("   pip install -r requirements.txt")
        sys.exit(1)
        print("   python setup_project.py")
        sys.exit(1)

if __name__ == "__main__":
    main()