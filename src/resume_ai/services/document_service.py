"""Service for handling document operations."""
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.document import Document as DocumentType

from resume_ai.utils import (
    ensure_extension,
    get_unique_filename,
    read_file as read_file_util,
    ensure_directory
)

class DocumentService:
    """Service for handling document operations."""
    
    @staticmethod
    def read_file(file_path: str) -> str:
        """Read content from a file.
        
        Args:
            file_path: Path to the file to read
            
        Returns:
            str: The file contents
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            IOError: If there's an error reading the file
        """
        return read_file_util(file_path)
    
    @classmethod
    def update_docx_objective(
        cls,
        source_path: str,
        output_path: str,
        new_objective: str,
        placeholder: str = '<objective_here>',
        font_name: str = 'Spectral',
        font_size: int = 10
    ) -> str:
        """
        Update the objective section in a DOCX file.
        
        Args:
            source_path: Path to the source DOCX file
            output_path: Path where to save the updated file
            new_objective: The objective text to insert
            placeholder: The placeholder text to replace
            font_name: Font name for the objective text
            font_size: Font size in points
            
        Returns:
            str: Path to the saved file
        """
        # Ensure output directory exists and get a unique filename
        output_path = Path(output_path)
        ensure_directory(output_path.parent)
        output_path = get_unique_filename(output_path)
        
        doc = Document(source_path)
        placeholder_found = False
        
        # Try to find and replace the placeholder
        for para in doc.paragraphs:
            if placeholder in para.text:
                para.clear()
                # Add 'OBJECTIVE' in bold
                # title_run = para.add_run('OBJECTIVE\n')
                # title_run.bold = True
                # title_run.font.name = font_name
                # title_run.font.size = Pt(font_size)
                # Add objective text with normal weight
                obj_run = para.add_run(new_objective)
                obj_run.font.name = font_name
                obj_run.font.size = Pt(font_size)
                placeholder_found = True
                break
        
        # If placeholder not found, add a new section at the beginning
        if not placeholder_found:
            para = doc.add_paragraph()
            # Add 'OBJECTIVE' in bold with specified font
            # title_run = para.add_run('OBJECTIVE\n')
            # title_run.bold = True
            # title_run.font.name = font_name
            # title_run.font.size = Pt(font_size)
            # Add objective text with normal weight
            obj_run = para.add_run(new_objective)
            obj_run.font.name = font_name
            obj_run.font.size = Pt(font_size)
        
        # Save the document
        doc.save(output_path)
        return output_path
